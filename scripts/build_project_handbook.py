from __future__ import annotations

import math
from datetime import date
from pathlib import Path
from textwrap import dedent

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "OpenTrial_Project_Handbook.docx"
ASSETS = ROOT / "docs" / "_handbook_assets"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 99, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BORDER = "D7DBE2"
INK = RGBColor(30, 34, 40)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in: list[float], indent_dxa: int = 120) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(int(w * 1440) for w in widths_in)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_in:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width_dxa = int(widths_in[idx] * 1440)
            cell.width = Inches(widths_in[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_border_bottom(paragraph, color="D7DBE2", size="8", space="6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    code = doc.styles.add_style("Code Block", 1)
    code.font.name = "Courier New"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Courier New")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Courier New")
    code.font.size = Pt(8.5)
    code.font.color.rgb = RGBColor(35, 39, 47)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.05

    caption = doc.styles.add_style("Figure Caption", 1)
    caption.font.name = "Calibri"
    caption.font.size = Pt(9.5)
    caption.font.color.rgb = MUTED
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    callout_style = doc.styles.add_style("Callout Text", 1)
    callout_style.font.name = "Calibri"
    callout_style.font.size = Pt(10.5)
    callout_style.paragraph_format.space_after = Pt(4)
    callout_style.paragraph_format.line_spacing = 1.25


def add_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("OpenTrial Project Handbook")
    set_run_font(run, size=9, color=MUTED)


def add_cover(doc: Document) -> None:
    section = doc.sections[0]
    add_footer(section)
    header = section.header.paragraphs[0]
    header.text = ""
    r = header.add_run("OpenTrial")
    set_run_font(r, size=9, color=MUTED, bold=True)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph_border_bottom(header, color="D7DBE2", size="4", space="2")

    doc.add_paragraph()
    doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = kicker.add_run("TECHNICAL OWNER'S FIELD GUIDE")
    set_run_font(kr, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("OpenTrial Project Handbook")
    set_run_font(tr, size=28, color=NAVY, bold=True)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run(
        "A didactic guide to the current codebase, design choices, computation path, "
        "and next build steps"
    )
    set_run_font(sr, size=13.5, color=MUTED)
    subtitle.paragraph_format.space_after = Pt(20)

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    rows = [
        ("Project", "OpenTrial - Bayesian clinical trial design engine"),
        ("Current slice", "Streamlit form -> seeded evidence -> prior -> simulation -> Markdown report"),
        ("Prepared for", "Reuben N Addison"),
        ("Generated", date.today().isoformat()),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=10, color=INK, bold=(cell is row.cells[0]))
    set_table_geometry(meta, [1.55, 4.75])

    doc.add_paragraph()
    callout = add_callout(
        doc,
        "Core idea",
        "The first milestone is deliberately small: make the engine produce a credible, "
        "inspectable report before adding Gemini orchestration or seven live data sources. "
        "This is how you keep the project understandable and ownable.",
    )
    callout.paragraph_format.space_after = Pt(28)
    doc.add_page_break()


def add_callout(doc: Document, label: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [6.3])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.style = "Callout Text"
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, size=10.5, color=DARK_BLUE, bold=True)
    r2 = p.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    return p


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def add_code_block(doc: Document, code: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [6.3])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F9FC")
    p = cell.paragraphs[0]
    p.style = "Code Block"
    p.paragraph_format.space_after = Pt(0)
    lines = code.rstrip("\n").splitlines()
    if not lines:
        lines = [""]
    for i, line in enumerate(lines):
        if i:
            p.add_run().add_break()
        run = p.add_run(line[:120])
        set_run_font(run, name="Courier New", size=8.3, color=RGBColor(30, 34, 40))


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE)
        for p in cell.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_run_font(run, size=9.5, color=NAVY, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=9.2, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph()


def font(size=20, bold=False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "",
        "/Library/Fonts/Arial.ttf",
    ]
    for path in candidates:
        if path and Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def draw_box(draw, xy, text, fill, outline="#6D7C91", text_fill="#0B2545", radius=18):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=2)
    x1, y1, x2, y2 = xy
    words = text.split()
    lines = []
    line = ""
    f = font(24, bold=True)
    for word in words:
        candidate = f"{line} {word}".strip()
        if draw.textlength(candidate, font=f) <= (x2 - x1 - 34):
            line = candidate
        else:
            lines.append(line)
            line = word
    if line:
        lines.append(line)
    total_h = len(lines) * 30
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        w = draw.textlength(line, font=f)
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=f, fill=text_fill)
        y += 30


def draw_arrow(draw, start, end, color="#2E74B5"):
    draw.line([start, end], fill=color, width=5)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 18
    for delta in (math.pi * 0.82, -math.pi * 0.82):
        x = end[0] + length * math.cos(angle + delta)
        y = end[1] + length * math.sin(angle + delta)
        draw.line([end, (x, y)], fill=color, width=5)


def make_diagram(name: str, title: str, boxes: list[tuple[str, tuple[int, int, int, int]]], arrows):
    ASSETS.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1400, 620), "#FFFFFF")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, 1400, 620), fill="#FFFFFF")
    draw.text((50, 34), title, font=font(32, bold=True), fill="#0B2545")
    draw.line((50, 84, 1350, 84), fill="#D7DBE2", width=3)
    for start, end in arrows:
        draw_arrow(draw, start, end)
    palette = ["#E8EEF5", "#F4F6F9", "#FFF7E6", "#EAF7EF", "#F7EEF8", "#EEF7FA"]
    for idx, (text, xy) in enumerate(boxes):
        draw_box(draw, xy, text, palette[idx % len(palette)])
    out = ASSETS / name
    img.save(out)
    return out


def build_diagrams():
    flow = make_diagram(
        "vertical_slice_flow.png",
        "Figure 1. First vertical slice: from input to report",
        [
            ("Streamlit form", (60, 190, 260, 300)),
            ("TrialDesignInput", (330, 190, 570, 300)),
            ("Demo evidence", (640, 110, 880, 220)),
            ("Build prior", (640, 290, 880, 400)),
            ("Simulate grid", (950, 190, 1170, 300)),
            ("Markdown report", (1220, 190, 1370, 300)),
        ],
        [
            ((260, 245), (330, 245)),
            ((570, 245), (640, 165)),
            ((570, 245), (640, 345)),
            ((880, 345), (950, 245)),
            ((880, 165), (950, 245)),
            ((1170, 245), (1220, 245)),
        ],
    )
    fallback = make_diagram(
        "fallback_live_pattern.png",
        "Figure 2. Mock now, live later: the CABSweb-inspired pattern",
        [
            ("config.py reads env", (70, 170, 320, 280)),
            ("Capability flags", (410, 170, 650, 280)),
            ("Seeded evidence path", (760, 95, 1050, 205)),
            ("Live API path", (760, 265, 1050, 375)),
            ("Same schema output", (1130, 170, 1340, 280)),
        ],
        [
            ((320, 225), (410, 225)),
            ((650, 225), (760, 150)),
            ((650, 225), (760, 320)),
            ((1050, 150), (1130, 225)),
            ((1050, 320), (1130, 225)),
        ],
    )
    compute = make_diagram(
        "computation_path.png",
        "Figure 3. Computation path and current statistical assumptions",
        [
            ("Effect estimates + SE", (70, 170, 360, 280)),
            ("Inverse-variance weights", (440, 170, 730, 280)),
            ("Normal prior mean + SD", (810, 170, 1090, 280)),
            ("One-sided z-test power", (440, 360, 730, 470)),
            ("Recommended N/arm", (810, 360, 1090, 470)),
        ],
        [
            ((360, 225), (440, 225)),
            ((730, 225), (810, 225)),
            ((590, 280), (590, 360)),
            ((730, 415), (810, 415)),
        ],
    )
    structure = make_diagram(
        "package_structure.png",
        "Figure 4. Module boundaries: each folder owns one concern",
        [
            ("app.py UI", (70, 170, 270, 280)),
            ("schemas.py contracts", (350, 170, 590, 280)),
            ("data/ evidence", (670, 95, 900, 205)),
            ("compute/ math", (670, 265, 900, 375)),
            ("report/ output", (980, 170, 1210, 280)),
            ("tests/ assurance", (980, 360, 1210, 470)),
        ],
        [
            ((270, 225), (350, 225)),
            ((590, 225), (670, 150)),
            ((590, 225), (670, 320)),
            ((900, 150), (980, 225)),
            ((900, 320), (980, 225)),
            ((1095, 280), (1095, 360)),
        ],
    )
    return [flow, fallback, compute, structure]


def read(path: str) -> str:
    return (ROOT / path).read_text()


def add_source_appendix(doc: Document, path: str) -> None:
    doc.add_heading(path, level=2)
    add_body_paragraph(
        doc,
        "Full current source listing. Read this after the chapter explanation so the code has context.",
    )
    add_code_block(doc, read(path))


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    diagrams = build_diagrams()

    doc = Document()
    configure_styles(doc)
    add_cover(doc)

    doc.add_heading("How To Use This Handbook", level=1)
    add_body_paragraph(
        doc,
        "This document is written as an owner manual for the current OpenTrial repository. "
        "It explains what exists today, why it was structured that way, how data moves "
        "through the code, how to run and test the app, and how to extend it without "
        "losing control of the design."
    )
    add_callout(
        doc,
        "Reading path",
        "Read Chapters 1-4 first for the mental model. Use Chapters 5-9 while coding. "
        "Use the appendices as the source-code reference for every file currently in the project.",
    )
    doc.add_heading("Static Table of Contents", level=2)
    toc_items = [
        "1. Project North Star",
        "2. Current Repository State",
        "3. Architectural Design",
        "4. Runtime Walkthrough",
        "5. File-by-File Explanation",
        "6. Statistical Computation",
        "7. Report Generation",
        "8. Running, Testing, and Debugging",
        "9. If Building From Scratch",
        "10. Next Build Steps",
        "Appendix A. Full Source Listings",
    ]
    for item in toc_items:
        add_bullet(doc, item)
    doc.add_page_break()

    doc.add_heading("1. Project North Star", level=1)
    add_body_paragraph(
        doc,
        "OpenTrial is being shaped as a computation engine for Bayesian clinical trial "
        "design. The final vision described in the MVP is ambitious: a Streamlit form "
        "collects trial design inputs, an agent gathers public evidence, the engine builds "
        "a Bayesian prior, simulation estimates operating characteristics, and the system "
        "returns a cited report."
    )
    add_body_paragraph(
        doc,
        "The important product decision is that the current project starts with the engine, "
        "not the agent. The agent can summarize, coordinate, and cite once the deterministic "
        "engine is trustworthy. This order keeps the math inspectable and prevents the LLM "
        "from becoming the place where hidden logic lives."
    )
    add_table(
        doc,
        ["Layer", "Current purpose", "Why it matters"],
        [
            ["Streamlit UI", "Collect structured trial inputs.", "Forces the problem into reproducible fields instead of chatty ambiguity."],
            ["Schemas", "Define shared contracts between modules.", "Prevents each function from inventing its own shape for evidence or design data."],
            ["Seeded evidence", "Provide deterministic demo records.", "Makes the app work with zero secrets and gives us known test data."],
            ["Computation", "Build prior and sample-size grid.", "Keeps numeric truth outside the LLM."],
            ["Report", "Turn results into a shareable Markdown artifact.", "Creates the first real deliverable users can inspect and download."],
        ],
        [1.25, 2.25, 2.7],
    )

    doc.add_heading("2. Current Repository State", level=1)
    add_body_paragraph(
        doc,
        "Before implementation, OpenTrial contained only planning documents. Step 1 added "
        "the first executable scaffold. The repository is still intentionally small; that is "
        "a strength at this stage because every file has one job."
    )
    add_code_block(
        doc,
        dedent(
            """
            OpenTrial/
              app.py
              pyproject.toml
              .env.example
              .gitignore
              README.md
              MVP.md
              VALUE.md
              src/opentrial/
                __init__.py
                config.py
                schemas.py
                data/demo_evidence.py
                integrations/registry.py
                compute/priors.py
                compute/simulation.py
                report/markdown.py
              tests/test_vertical_slice.py
            """
        ).strip(),
    )
    add_body_paragraph(
        doc,
        "The code follows a src-layout package structure. That means importable project "
        "code lives under src/opentrial rather than beside app.py. This avoids accidental "
        "imports from the working directory and makes packaging, tests, and later deployment cleaner."
    )

    doc.add_heading("3. Architectural Design", level=1)
    doc.add_picture(str(diagrams[0]), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 1. The current app is one complete vertical slice.")
    cap.style = "Figure Caption"
    add_body_paragraph(
        doc,
        "The design is deliberately modular. The UI should not know how priors are computed. "
        "The computation layer should not know how Streamlit displays charts. The report layer "
        "should not fetch data. This keeps each part replaceable."
    )
    doc.add_picture(str(diagrams[1]), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 2. The fallback/live pattern was borrowed from CABSweb.")
    cap.style = "Figure Caption"
    add_body_paragraph(
        doc,
        "The CABSweb pattern worth copying is not its frontend stack; it is the capability-flag "
        "idea. OpenTrial should run end-to-end with no secrets, using deterministic seed data. "
        "When environment variables are present, the same function boundaries can switch to live adapters."
    )

    doc.add_heading("4. Runtime Walkthrough", level=1)
    add_body_paragraph(
        doc,
        "When you press Generate design report, app.py performs a short deterministic workflow. "
        "This is the project heartbeat. If this sequence stays clear, the project remains ownable."
    )
    for step in [
        "Collect values from the Streamlit form: indication, endpoint, target effect, alpha, desired power, and maximum sample size per arm.",
        "Create a TrialDesignInput dataclass instance. This freezes the UI values into one typed object.",
        "Load the seeded T2D/HbA1c evidence pack from demo_evidence.py.",
        "Pass evidence records into build_prior(), which calculates an inverse-variance normal prior.",
        "Pass the design and prior into simulate_design_grid(), which estimates power at sample-size increments.",
        "Choose the first grid point that reaches desired power with recommend_sample_size().",
        "Render a Markdown report and expose it in the UI with a download button.",
    ]:
        add_numbered(doc, step)
    add_callout(
        doc,
        "Ownership test",
        "You should be able to point to any screen value and trace it to a dataclass field, a function input, "
        "a line in the report, or a test assertion. That traceability is the design standard.",
    )

    doc.add_heading("5. File-by-File Explanation", level=1)
    add_table(
        doc,
        ["File", "Responsibility", "Design reason"],
        [
            ["app.py", "Streamlit UI and orchestration.", "Thin UI layer; calls domain functions but does not own math."],
            ["pyproject.toml", "Package metadata, dependencies, pytest config.", "Makes install/test commands reproducible."],
            [".env.example", "Documents optional secrets and flags.", "Shows how live mode will be enabled later."],
            ["config.py", "Environment parsing and capability flags.", "Centralizes runtime decisions."],
            ["schemas.py", "Dataclasses used across modules.", "Creates stable contracts before APIs are added."],
            ["demo_evidence.py", "Seeded T2D/HbA1c evidence.", "Allows zero-secret demos and tests."],
            ["registry.py", "Integration status display.", "Makes live/mock state visible to the user."],
            ["priors.py", "Evidence-derived prior calculation.", "Keeps prior math isolated and testable."],
            ["simulation.py", "Power and posterior-success grid.", "Keeps sample-size behavior inspectable."],
            ["markdown.py", "Report string rendering.", "Separates computation from deliverable formatting."],
            ["test_vertical_slice.py", "End-to-end smoke test.", "Protects the first working project heartbeat."],
        ],
        [1.55, 2.15, 2.5],
    )

    doc.add_heading("5.1 app.py", level=2)
    add_body_paragraph(
        doc,
        "app.py is the only Streamlit-specific file. It imports the domain modules, creates widgets, "
        "builds a TrialDesignInput, and displays results. The code intentionally avoids putting "
        "statistical formulas in the UI file."
    )
    add_code_block(
        doc,
        dedent(
            """
            design = TrialDesignInput(
                indication=indication,
                endpoint=endpoint,
                target_effect=float(target_effect),
                alpha=float(alpha),
                desired_power=float(desired_power),
                max_n_per_arm=int(max_n),
            )
            """
        ).strip(),
    )
    add_body_paragraph(
        doc,
        "This block is the boundary between interface and engine. Everything below it should "
        "accept structured project objects rather than raw widget values."
    )

    doc.add_heading("5.2 schemas.py", level=2)
    add_body_paragraph(
        doc,
        "Schemas are the nouns of the project. They answer: what is a trial design input, what is an "
        "evidence record, what is a prior, and what is a simulated design point?"
    )
    add_table(
        doc,
        ["Dataclass", "Meaning"],
        [
            ["TrialDesignInput", "The structured question the user is asking."],
            ["EvidenceRecord", "One normalized piece of evidence from CT.gov, PubMed, or a future source."],
            ["PriorSummary", "The mathematical prior produced from the evidence."],
            ["DesignPoint", "One sample-size point with estimated operating characteristics."],
            ["IntegrationStatus", "A UI-visible service status row."],
        ],
        [1.7, 4.5],
    )

    doc.add_heading("5.3 config.py and registry.py", level=2)
    add_body_paragraph(
        doc,
        "config.py reads environment variables once and exposes capability flags. registry.py turns "
        "those flags into human-readable integration states for the sidebar. This keeps future live "
        "API logic from leaking into the UI."
    )
    add_code_block(
        doc,
        dedent(
            """
            @property
            def pubmed_enabled(self) -> bool:
                return self.use_live_apis and bool(self.ncbi_email)
            """
        ).strip(),
    )
    add_body_paragraph(
        doc,
        "Notice the two-condition design: PubMed is considered live only when live APIs are enabled "
        "and NCBI email is configured. This makes accidental network mode less likely."
    )

    doc.add_heading("5.4 demo_evidence.py", level=2)
    add_body_paragraph(
        doc,
        "The demo evidence pack is not pretending to be real extracted evidence. It is synthetic, "
        "normalized, and clearly labeled. Its purpose is to let the computation and report path work "
        "before we build fragile network integrations."
    )
    add_table(
        doc,
        ["Record", "Effect", "SE", "N", "Role in the demo"],
        [
            ["CT.gov SGLT2-like precedent", "0.48", "0.12", "256", "Trial precedent"],
            ["CT.gov GLP-1-like precedent", "0.62", "0.15", "188", "Trial precedent"],
            ["PubMed incretin meta-analysis", "0.55", "0.10", "1140", "Literature effect estimate"],
            ["PubMed SGLT2 review", "0.43", "0.09", "980", "Literature effect estimate"],
        ],
        [1.7, 0.7, 0.55, 0.55, 2.7],
    )

    doc.add_heading("6. Statistical Computation", level=1)
    doc.add_picture(str(diagrams[2]), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 3. Current computation path and assumptions.")
    cap.style = "Figure Caption"
    add_body_paragraph(
        doc,
        "The current statistical engine is intentionally simple and transparent. It is not the final "
        "Bayesian adaptive design promised by the MVP, but it creates a correct place for that future "
        "work to live."
    )
    doc.add_heading("6.1 Building the prior", level=2)
    add_body_paragraph(
        doc,
        "The prior is built using inverse-variance weighting. Evidence records with smaller standard "
        "errors receive larger weights. Conceptually:"
    )
    add_code_block(
        doc,
        dedent(
            """
            weight_i = 1 / standard_error_i^2
            prior_mean = sum(weight_i * effect_i) / sum(weight_i)
            fixed_sd = sqrt(1 / sum(weight_i))
            prior_sd = sqrt(fixed_sd^2 + heterogeneity^2)
            """
        ).strip(),
    )
    add_body_paragraph(
        doc,
        "The heterogeneity inflation is a practical guardrail. If prior studies disagree, the prior "
        "should become wider rather than overconfident."
    )
    doc.add_heading("6.2 Simulating sample-size grid", level=2)
    add_body_paragraph(
        doc,
        "simulation.py currently approximates power with a one-sided z-test for a continuous two-arm endpoint. "
        "It assumes a standardized endpoint standard deviation of 1.0. That is a placeholder assumption, "
        "but it is explicit and easy to replace later."
    )
    add_code_block(
        doc,
        dedent(
            """
            standard_error = sqrt((2 * endpoint_sd^2) / n_per_arm)
            z_alpha = normal_quantile(1 - alpha)
            z_effect = effect / standard_error
            power = 1 - normal_cdf(z_alpha - z_effect)
            """
        ).strip(),
    )
    add_callout(
        doc,
        "Current limitation",
        "This is not yet a full PyMC group-sequential adaptive design. It is a transparent first model "
        "that gives the app a mathematically coherent behavior while we build the data and report shell.",
    )

    doc.add_heading("7. Report Generation", level=1)
    add_body_paragraph(
        doc,
        "The first deliverable format is Markdown. Markdown is intentionally chosen before PDF because "
        "it is easy to inspect, version, test, and download. PDF export can come later once the report "
        "content stabilizes."
    )
    add_table(
        doc,
        ["Report section", "Generated from"],
        [
            ["Design Question", "TrialDesignInput"],
            ["Evidence-Derived Prior", "PriorSummary"],
            ["Recommendation", "DesignPoint selected by recommend_sample_size()"],
            ["Operating Characteristics", "List of DesignPoint values"],
            ["Evidence Provenance", "EvidenceRecord list"],
            ["Notes", "Known limitations and next live-adapter step"],
        ],
        [2.1, 4.1],
    )
    add_body_paragraph(
        doc,
        "The key design rule for reports is provenance. Any number in the report should be traceable "
        "back to either user input, evidence records, or a computation function."
    )

    doc.add_heading("8. Running, Testing, and Debugging", level=1)
    add_heading = doc.add_heading
    add_heading("8.1 Install", level=2)
    add_code_block(
        doc,
        dedent(
            """
            cd /Users/reubenaddison/Documents/Cabs/OpenTrial
            python3 -m venv .venv
            .venv/bin/python -m pip install -e '.[dev]'
            """
        ).strip(),
    )
    add_heading("8.2 Run the app", level=2)
    add_code_block(doc, ".venv/bin/streamlit run app.py --server.port 8501")
    add_body_paragraph(
        doc,
        "Open http://localhost:8501. Use the default T2D/HbA1c fields, press Generate design report, "
        "and confirm that metrics, the line chart, Markdown report, and download button appear."
    )
    add_heading("8.3 Run tests", level=2)
    add_code_block(doc, ".venv/bin/python -m pytest -q")
    add_body_paragraph(
        doc,
        "The current test is intentionally an end-to-end smoke test for the first vertical slice. "
        "It proves that seeded evidence can become a prior, a simulation grid, a recommendation, and a report."
    )
    add_heading("8.4 Debugging guide", level=2)
    add_table(
        doc,
        ["Symptom", "Likely cause", "Where to look"],
        [
            ["Streamlit missing", "Dependencies not installed in .venv.", "pyproject.toml and install command"],
            ["No recommendation", "Desired power not reached before max_n_per_arm.", "simulation.py and UI max N"],
            ["Prior mean seems wrong", "Evidence effect or SE values changed.", "demo_evidence.py and priors.py"],
            ["Sidebar says mock", "Expected until live flags and keys are set.", "config.py and .env.example"],
            ["Report missing section", "Markdown renderer changed.", "report/markdown.py and test_vertical_slice.py"],
        ],
        [1.55, 2.15, 2.5],
    )

    doc.add_heading("9. If Building From Scratch", level=1)
    add_body_paragraph(
        doc,
        "If you were starting OpenTrial from an empty folder, do not begin with Gemini, "
        "ClinicalTrials.gov, PubMed, or Streamlit polish. Begin by defining the objects "
        "that must move through the system. The safest build order is: nouns first, fake "
        "data second, computation third, report fourth, interface fifth, live integrations sixth, "
        "agent orchestration last."
    )
    add_callout(
        doc,
        "Unambiguous rule",
        "A module is ready only when you can state what it receives, what it returns, what it must not know, "
        "and how it is tested. If you cannot answer those four questions, the boundary is not clear yet.",
    )

    doc.add_heading("9.1 The exact build order", level=2)
    add_body_paragraph(
        doc,
        "The following order is intentionally conservative. It prevents you from wiring a beautiful UI "
        "around unclear data, or wiring live APIs into code that has no stable internal contracts."
    )
    for step in [
        "Create the package shell: pyproject.toml, src/opentrial/__init__.py, and tests/. This gives Python a real project shape.",
        "Write schemas.py. Define TrialDesignInput, EvidenceRecord, PriorSummary, DesignPoint, and IntegrationStatus before writing business logic.",
        "Write data/demo_evidence.py. Add a small T2D/HbA1c evidence pack that returns EvidenceRecord objects.",
        "Write compute/priors.py. Make one function that turns evidence into a PriorSummary.",
        "Write compute/simulation.py. Make one function that turns TrialDesignInput plus PriorSummary into a list of DesignPoint objects.",
        "Write report/markdown.py. Make the first artifact plain Markdown so the output is easy to inspect.",
        "Write tests/test_vertical_slice.py. Prove the whole path works without Streamlit, Gemini, or network calls.",
        "Write app.py. Let Streamlit call the already-working engine. Do not put math in the UI.",
        "Only after that, add tools/clinicaltrials.py and tools/pubmed.py as live evidence adapters.",
        "Only after live evidence works, add Gemini as an orchestrator and report-writer, not as the math engine.",
    ]:
        add_numbered(doc, step)

    doc.add_heading("9.2 The module boundary rules", level=2)
    add_body_paragraph(
        doc,
        "Each module exists because it owns one category of decision. If a module starts making decisions "
        "outside its category, the project becomes hard to reason about. Use this table as the boundary contract."
    )
    add_table(
        doc,
        ["Module", "Allowed to know", "Must not know"],
        [
            ["app.py", "Streamlit widgets, user clicks, display layout.", "Statistical formulas, raw API response shapes, report internals."],
            ["schemas.py", "Field names, object shapes, type contracts.", "How evidence is fetched, displayed, or statistically modeled."],
            ["data/", "Seeded evidence records and demo fixtures.", "Streamlit, Gemini, PDF formatting, or power formulas."],
            ["tools/", "External APIs and response normalization.", "UI layout, report prose, or downstream sample-size decisions."],
            ["compute/", "Priors, power, posterior summaries, recommendations.", "Where evidence came from or how results are displayed."],
            ["report/", "How structured results become Markdown or later PDF.", "How priors are calculated or APIs are called."],
            ["tests/", "Expected behavior and regression protection.", "Production secrets or live-only assumptions."],
        ],
        [1.25, 2.55, 2.4],
    )

    doc.add_heading("9.3 Why schemas come first", level=2)
    add_body_paragraph(
        doc,
        "Schemas are the vocabulary of the engine. Without them, every new function invents its own version "
        "of a trial, an evidence item, or a result. That seems faster for the first hour, then becomes expensive "
        "when the UI, computation layer, report layer, and tests all disagree about field names."
    )
    add_code_block(
        doc,
        dedent(
            """
            # Good first question:
            What exact object does the next function need?

            # Bad first question:
            Which API or UI widget should I wire up first?
            """
        ).strip(),
    )
    add_body_paragraph(
        doc,
        "For OpenTrial, the first critical object is TrialDesignInput because it defines the user's design "
        "question. The second is EvidenceRecord because it defines how every database must eventually speak "
        "to the engine. Once these shapes exist, everything else can be built against them."
    )

    doc.add_heading("9.4 What not to build first", level=2)
    add_table(
        doc,
        ["Tempting first move", "Why it is risky", "Better first move"],
        [
            ["Start with Gemini", "The model may produce plausible prose before the engine is correct.", "Build deterministic functions, then let Gemini orchestrate them."],
            ["Start with all seven APIs", "You spend days handling network quirks before knowing your internal data shape.", "Start with seeded EvidenceRecord objects."],
            ["Start with PDF export", "You polish a document format before the content is stable.", "Start with Markdown report generation."],
            ["Start with a complex PyMC model", "The math becomes harder to debug before the data path is proven.", "Start with a transparent normal approximation and test it."],
            ["Start with a large UI", "Widgets multiply before the underlying workflow is real.", "Start with one vertical slice and default values."],
        ],
        [1.55, 2.45, 2.2],
    )

    doc.add_heading("9.5 Definition of done for each stage", level=2)
    add_body_paragraph(
        doc,
        "Use these acceptance checks to keep the build honest. A stage is not done because the code exists; "
        "it is done when the next stage can rely on it without guessing."
    )
    add_table(
        doc,
        ["Stage", "Done means"],
        [
            ["Schemas", "Every downstream object has a clear dataclass and all field names are stable enough for tests."],
            ["Seed data", "The project can run with no secrets, no network, and predictable evidence values."],
            ["Prior computation", "EvidenceRecord list reliably produces a PriorSummary, including a fallback if no usable evidence exists."],
            ["Simulation", "TrialDesignInput plus PriorSummary produces a sample-size grid and a recommendation when power is reached."],
            ["Report", "The report includes design inputs, prior, recommendation, operating characteristics, provenance, and limitations."],
            ["UI", "A non-coder can generate and download the report without touching Python."],
            ["Live adapter", "The adapter returns the same internal schema as seed data and has a cached-fixture test."],
            ["Agent", "Gemini calls existing tools and writes synthesis; it does not invent hidden calculations."],
        ],
        [1.55, 4.65],
    )

    doc.add_heading("9.6 Improvements added in this revision", level=2)
    add_body_paragraph(
        doc,
        "This revision of the handbook adds explicit build instructions because the previous version explained "
        "the current codebase well but left some tacit engineering judgment unstated. The improved guidance makes "
        "the architecture easier to reproduce from scratch."
    )
    for improvement in [
        "Added a strict build order so you know what to create first and what to postpone.",
        "Added module boundary rules so each file has a clear responsibility and a clear list of things it must not do.",
        "Added a list of tempting but risky first moves, such as starting with Gemini or all seven APIs.",
        "Added definitions of done so progress can be judged by working behavior, not just file creation.",
        "Made the agent role explicit: Gemini should orchestrate and synthesize after the engine works.",
    ]:
        add_bullet(doc, improvement)

    doc.add_heading("10. Next Build Steps", level=1)
    add_body_paragraph(
        doc,
        "The next engineering move is to add the ClinicalTrials.gov wrapper while preserving the seeded "
        "fallback. Do not replace the demo path; live adapters should return the same EvidenceRecord shape."
    )
    for step in [
        "Create src/opentrial/tools/clinicaltrials.py with a function such as get_trials_ct_gov(indication, n=30).",
        "Normalize the live response into EvidenceRecord objects, even if early records use placeholder effect fields.",
        "Add config logic so OPENTRIAL_USE_LIVE_APIS=true switches the evidence source from seed data to live CT.gov.",
        "Add tests with cached fixture JSON so CI and demos still work offline.",
        "Only after CT.gov is stable, add PubMed E-utilities and normalize literature effect estimates.",
    ]:
        add_numbered(doc, step)
    doc.add_picture(str(diagrams[3]), width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph("Figure 4. Folder boundaries should stay stable as live wrappers are added.")
    cap.style = "Figure Caption"

    doc.add_page_break()
    doc.add_heading("Appendix A. Full Source Listings", level=1)
    add_body_paragraph(
        doc,
        "This appendix includes every project source file added in the first scaffold. "
        "Use it as a study copy: read a chapter, then read the corresponding source."
    )
    for path in [
        "pyproject.toml",
        ".env.example",
        ".gitignore",
        "app.py",
        "src/opentrial/__init__.py",
        "src/opentrial/config.py",
        "src/opentrial/schemas.py",
        "src/opentrial/data/__init__.py",
        "src/opentrial/data/demo_evidence.py",
        "src/opentrial/integrations/__init__.py",
        "src/opentrial/integrations/registry.py",
        "src/opentrial/compute/__init__.py",
        "src/opentrial/compute/priors.py",
        "src/opentrial/compute/simulation.py",
        "src/opentrial/report/__init__.py",
        "src/opentrial/report/markdown.py",
        "tests/test_vertical_slice.py",
    ]:
        add_source_appendix(doc, path)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
